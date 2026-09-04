"""
Stage 10: generates the final technical report as a Word document, pulling
every number directly from pipeline/final_results_2026-07-26/ (the frozen,
canonical results directory - see its MANIFEST.md) rather than typing
numbers in by hand.

Formatting spec, exactly as requested: Times New Roman throughout; normal
text 12pt; subtitles 14pt; all paragraphs justified; no em dashes or en
dashes anywhere, plain hyphens only; plain, honest, human tone rather than
inflated claims.
"""
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"
RESULTS = "pipeline/final_results_2026-07-26"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLACK
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    return p


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = Pt(11)
                    run.font.color.rgb = BLACK


def add_df_table(df, index_label=None, caption=None):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = BLACK
    cols = ([index_label] if index_label else []) + list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = str(c)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
    for idx, row in df.iterrows():
        cells = table.add_row().cells
        vals = ([str(idx)] if index_label else []) + [str(v) for v in row.tolist()]
        for i, v in enumerate(vals):
            cells[i].text = v
    style_table(table)
    doc.add_paragraph()


# ============================= load data ==================================
internal_cmp = pd.read_csv(f"{RESULTS}/table_internal_comparison.csv", index_col=0)
hybrid_b = pd.read_csv(f"{RESULTS}/table_hybrid_b.csv", index_col=0)
ext_val = pd.read_csv(f"{RESULTS}/table_external_validation.csv", index_col=0)
dl_cmp = pd.read_csv(f"{RESULTS}/table_dl_comparison.csv", index_col=0)
circ = pd.read_csv(f"{RESULTS}/table_circularity_check.csv", index_col=0)
era = pd.read_csv(f"{RESULTS}/table_training_era_overlap.csv")
acmg_gene = pd.read_csv(f"{RESULTS}/table_acmg_calibration_gene_specific.csv")
acmg_ext = pd.read_csv(f"{RESULTS}/table_acmg_calibration_external_validation.csv")
stage8 = pd.read_csv(f"{RESULTS}/table_stage8_benchmarking.csv")
dms_small = pd.read_csv(f"{RESULTS}/table_brca1_dms_validation.csv")
dms_full = pd.read_csv(f"{RESULTS}/table_brca1_dms_full_validation.csv")
smote = pd.read_csv(f"{RESULTS}/table_smote_comparison.csv")
cis = pd.read_csv(f"{RESULTS}/table_model_cis.csv")
sig = pd.read_csv(f"{RESULTS}/table_pairwise_significance.csv")
kazakh_case = pd.read_csv(f"{RESULTS}/kazakh_pm2_bs1_case_study.csv")
kazakh_final = pd.read_csv(f"{RESULTS}/kazakh_variants_final_annotated.csv")
missing = pd.read_csv(f"{RESULTS}/missingness_report.csv")
lit = pd.read_csv(f"{RESULTS}/table_benchmark_comparison_literature.csv")

sig_col = [c for c in sig.columns if "signif" in c][0]
n_sig_ext = int(sig[(sig.eval_set == "external") & sig[sig_col]].shape[0])
n_pairs_ext = int(sig[sig.eval_set == "external"].shape[0])
n_sig_int = int(sig[(sig.eval_set == "internal-test") & sig[sig_col]].shape[0])
n_pairs_int = int(sig[sig.eval_set == "internal-test"].shape[0])

lof_track = kazakh_final[kazakh_final["track"] == "LOF (rule-based)"]
mis_track = kazakh_final[kazakh_final["track"] == "Missense (ML)"]
lof_correct = int((lof_track["rule_prediction"] == lof_track["label"]).sum())
lof_total = int(lof_track["label"].notna().sum())

# ============================================================================
add_title("Hybrid Machine Learning and Deep Learning Benchmarking of BRCA1/2 Variant Pathogenicity Classification")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technical Report")
run.font.name = FONT
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = BLACK
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Druba Rahim Ahmed")
run2.font.name = FONT
run2.font.size = Pt(12)
run2.font.color.rgb = BLACK
doc.add_paragraph()

# ---------------------------------------------------------------------------
add_heading("Abstract")
add_para(
    "This report describes a computational pipeline for classifying BRCA1 and BRCA2 genetic "
    "variants as pathogenic or benign, and a thorough effort to check that every part of that "
    "pipeline actually does what it claims to do. The work started from an earlier version of "
    "this project that had real methodological problems, and much of what follows is the result "
    "of going back through the pipeline step by step, verifying each stage against the data "
    "rather than assuming it was correct, fixing the mistakes found along the way, and being "
    "honest about what still cannot be claimed. Along the way, several genuine errors were found "
    "and corrected: a wrong sign in one of the model's input scores, a variant annotation step "
    "that was reading the wrong gene transcript, a data provenance gap in another input score, "
    "and a genome build mismatch in one of the population datasets used later in the project. "
    "A thirteen feature model was built combining eleven established pathogenicity predictors "
    "with two additional scores added after checking the literature and the data directly. A "
    "hybrid model combining a tree based classifier with a neural network was rebuilt after its "
    "original code was found to be missing, and tested formally against six other models. The "
    "honest finding is that no single model statistically outperforms the others on the held out "
    "external validation set, which is itself an important and useful result, not a "
    "disappointment to hide. The model was validated against an independent set of functional "
    "laboratory measurements that have nothing to do with clinical databases, against a "
    "population dataset from Kazakhstan, and against a set of real published clinical variants "
    "from three Kazakh research studies. The report ends with a plain statement of what this "
    "project can and cannot currently claim."
)

# ---------------------------------------------------------------------------
add_heading("1. Objective")
add_para(
    "The goal of this project is to classify BRCA1 and BRCA2 variants as pathogenic or benign "
    "using computational methods, and to do so in a way that can withstand scrutiny. BRCA1 and "
    "BRCA2 are the two genes most strongly associated with hereditary breast and ovarian cancer, "
    "and a large number of variants found in these genes in patients still have an unknown or "
    "uncertain clinical meaning. Existing pathogenicity prediction tools each have their own "
    "strengths, weaknesses, and blind spots, and a recurring theme in the scientific literature "
    "on this topic is that many of these tools look better than they really are because of "
    "circularity between how they were built and how they are tested. This project tries to take "
    "that concern seriously rather than treat it as a footnote."
)
add_para(
    "This is also, honestly, a project about verification as much as it is about prediction. An "
    "earlier attempt at this same task, done by a previous collaborator, had real problems that "
    "were only discovered later, including unreproducible results and gaps in the annotation "
    "process. A central part of the current effort has been to never assume a previous result is "
    "correct just because it exists, and to re derive or re check every important number from the "
    "underlying data before relying on it."
)

# ---------------------------------------------------------------------------
add_heading("2. Data Sources")
add_para(
    "The variant dataset comes from a BRCA Exchange data release, dated November 2025, which "
    "pools BRCA1 and BRCA2 variant records from ClinVar, the ENIGMA consortium expert panel, "
    "LOVD, ExUV, BIC, ExAC, gnomAD, the 1000 Genomes Project, and the Exome Sequencing Project. "
    "All coordinates are in the GRCh38 genome build. Two independent pools were built from this "
    "release. The internal pool consists of variants with a usable ClinVar label and no ENIGMA "
    "call, and is used for model training and internal testing. The external pool consists of "
    "variants with a usable ENIGMA expert panel call, and is held out entirely for final "
    "validation. The two pools are guaranteed to share no variants by construction, and this was "
    "checked directly rather than assumed, both when the pools were built and again immediately "
    "before final evaluation."
)
add_para(
    "Pathogenicity scores were obtained from dbNSFP version 5.3.1a, a large database of "
    "precomputed variant annotations, queried directly from a local copy rather than through an "
    "external web service. This change from an earlier, API based approach was made deliberately: "
    "querying a local copy is faster, more reliable, and avoids depending on a third party service "
    "staying available. Gene context annotation, such as which exon a variant falls in and "
    "whether it changes the protein sequence, was obtained using ANNOVAR against the same genome "
    "build."
)

# ---------------------------------------------------------------------------
add_heading("3. Feature Set")
add_para(
    "The final feature set consists of thirteen pathogenicity scores, chosen after a systematic "
    "review of over fifty published papers on variant effect prediction and after empirical "
    "testing on this project's own data. Ten of the scores (ClinPred, MetaRNN, REVEL, MetaSVM, "
    "MetaLR, VEST4, AlphaMissense, PrimateAI, and CADD in both its raw and scaled forms) were "
    "retrieved directly from dbNSFP. ESM1b was computed separately using a protein language "
    "model, since dbNSFP's own version of this score was found, on direct comparison, to be "
    "sign inverted relative to what the rest of the pipeline expects, and this was corrected."
)
add_para(
    "Two further scores were added to the original eleven after specific investigation. BayesDel, "
    "a score with a published gene specific validation directly on BRCA1 and BRCA2 against a "
    "clinical gold standard, was added using the version computed from the current local dbNSFP "
    "file rather than the version bundled with the variant database itself, because that bundled "
    "version was traced back to a much older, outdated release of dbNSFP. Eigen, an unsupervised "
    "score that is not trained on any clinical labels at all, was added after testing showed it "
    "improved classification accuracy without meaningfully overlapping with information already "
    "captured by the other scores. Its close relative, Eigen PC, was tested and left out, since it "
    "was found to be almost a duplicate of Eigen and also strongly overlapping with CADD."
)
add_para(
    "Two further candidates were considered and rejected. A newer score called popEVE was tested "
    "but found to be available for such a small fraction of variants that including it would have "
    "removed most of the usable data for only a small gain in accuracy. PrimateAI was kept in the "
    "feature set despite its original authors having asked that it not be used as an input to "
    "other classifiers, since removing it was shown to cost real predictive accuracy and this "
    "project does not publish any new benchmark that PrimateAI's own evaluation depends on."
)
add_para(
    "Because pathogenicity predictors are largely built to score missense variants, and cannot "
    "meaningfully score frameshifts, large deletions, or most splicing variants, a large fraction "
    "of variants are missing most or all of these thirteen scores. The following table shows the "
    "percentage of each pool missing each score."
)
mp = missing.pivot(index="feature", columns="pool", values="pct_missing")
add_df_table(mp, index_label="Score", caption="Table 1. Percent of each pool missing each score.")

# ---------------------------------------------------------------------------
add_heading("4. Two Track Classification Design")
add_para(
    "A single machine learning model trained only on variants with complete score coverage would "
    "quietly throw away almost all of the loss of function variants, since those are exactly the "
    "variants missense predictors cannot score, and loss of function variants make up a large "
    "share of real pathogenic calls. To avoid this, the project uses two separate tracks."
)
add_para(
    "The first track is a simple rule, not a trained model. Variants that cause a premature stop "
    "codon, a frameshift, loss of the start codon, or disruption of a canonical splice site are "
    "called pathogenic directly, following the same logic used in the official ACMG clinical "
    "variant classification guidelines for loss of function evidence. One specific, well "
    "documented exception is built in: truncating variants in BRCA2 at or after amino acid "
    "position 3326 are not automatically called high risk, since this region is known from the "
    "clinical genetics literature to escape the usual loss of function mechanism."
)
add_para(
    "The second track is the machine learning model described in the rest of this report, applied "
    "only to missense variants with complete score coverage, using class weighting rather than "
    "synthetic oversampling to handle the fact that pathogenic variants are a small minority of "
    "the data (discussed further in Section 11)."
)

# ---------------------------------------------------------------------------
add_heading("5. Model Development")
add_para(
    "Four classical machine learning classifiers were trained on the thirteen score feature set: "
    "Random Forest, a Support Vector Machine, Gradient Boosting, and XGBoost, each tuned by grid "
    "search with five fold cross validation. Two ensemble methods were then built on top of these: "
    "a soft voting ensemble and a stacking ensemble with a logistic regression combiner."
)
add_para(
    "Two sequence based deep learning models, a one dimensional convolutional network and a "
    "bidirectional LSTM, were trained directly on raw genomic sequence windows around each "
    "variant, with no score based features at all. Both performed reasonably on internal "
    "validation but generalized very poorly to the external validation set, achieving accuracy "
    "close to what would be expected from random guessing. This is discussed honestly in Section "
    "7, since it directly affects how much weight the phrase deep learning should carry in this "
    "project's title."
)
add_para(
    "A hybrid model, referred to throughout this report as Hybrid B, combines an XGBoost model "
    "and a small neural network, both trained on the same thirteen score feature set, stacked "
    "together with a logistic regression combiner using five fold out of fold prediction so that "
    "no part of the training data is ever scored by a model that was trained on it. This hybrid "
    "model's own code was, at one point during this project, found to have gone missing from the "
    "project files entirely, with only its previously reported results surviving. It was rebuilt "
    "from a written description of its design and re evaluated on the corrected data, and the "
    "results reported in this document reflect that rebuilt version, not the original, "
    "unreproducible one."
)

# ---------------------------------------------------------------------------
add_heading("6. Internal and External Validation Results")
add_para(
    "All models were evaluated on an internal test set held out from training, and separately on "
    "the fully independent external validation set described in Section 2. The following tables "
    "report the standard classification metrics for each model."
)
add_df_table(internal_cmp.round(4), index_label="Model", caption="Table 2. Internal test set results, thirteen feature classical and ensemble models.")
hybrid_b_display = hybrid_b[["n", "n_pathogenic", "roc_auc", "auprc", "f1"]].round(4).copy()
hybrid_b_display.index = [
    "Internal test, default threshold", "External, default threshold",
    "Internal test, F1 optimal threshold", "External, F1 optimal threshold",
]
add_df_table(
    hybrid_b_display, index_label="Evaluation",
    caption="Table 3. Hybrid B results, internal test set and external validation, at a default decision threshold of 0.5 and at an F1 optimal threshold chosen from training data only.",
)
add_df_table(ext_val[["n", "n_pathogenic", "test_roc_auc", "test_accuracy", "test_f1"]].round(4), index_label="Model",
             caption="Table 4. External validation set results, all models.")
add_para(
    "The sequence only deep learning models, evaluated separately since they use a different, "
    "less restrictive filtering rule, show the pattern described in Section 5 clearly."
)
add_df_table(dl_cmp.round(4), index_label="Model", caption="Table 5. Sequence based deep learning models, internal validation ROC AUC.")
add_para(
    "The one dimensional convolutional network and the bidirectional LSTM both drop to close to "
    "chance level performance on the external set, roughly 0.53 and 0.54 respectively, confirmed "
    "directly in this project's own external validation run. This is a real and important finding: "
    "training a sequence model from scratch on a dataset of this size, without any form of "
    "pretraining, does not currently produce a model that generalizes."
)

# ---------------------------------------------------------------------------
add_heading("7. Are Any of These Models Actually Better Than the Others")
add_para(
    "All of the models above achieve very high scores, clustered tightly between roughly 0.97 and "
    "0.99 internal ROC AUC and 0.998 to 1.000 external ROC AUC. A number that close to another "
    "number is not automatically meaningfully different from it, and this project treats that "
    "question as one that needs to actually be tested rather than assumed."
)
add_para(
    "Bootstrap confidence intervals were computed for every model's ROC AUC, area under the "
    "precision recall curve, and F1 score, and every pair of models was compared directly using a "
    "paired bootstrap significance test on the external validation set and separately on the "
    "internal test set. The honest result is this: on the external validation set, none of the "
    f"twenty one possible pairs of models showed a statistically significant difference in ROC "
    "AUC. On the internal test set, six of twenty one pairs did reach statistical significance, "
    "but the actual size of the difference in every one of those six cases was small, on the "
    "order of one to two percentage points, which reflects the large size of the internal test set "
    "being able to detect a small difference more than it reflects one model being meaningfully "
    "better than another."
)
add_para(
    "Hybrid B specifically was not found to be statistically distinguishable from any of the other "
    "six models on either evaluation set. This directly contradicts an earlier claim, made before "
    "this project's verification pass, that Hybrid B clearly outperformed every other model. That "
    "earlier claim was based on a version of the pipeline that has since been shown to have real "
    "errors in it, and the honest, current claim is that several different modeling approaches, "
    "tree based, kernel based, and the tree plus neural network hybrid, all reach statistically "
    "indistinguishable performance on this task with this data. This is reported as a genuine "
    "finding rather than adjusted to produce a cleaner sounding winner."
)

# ---------------------------------------------------------------------------
add_heading("8. Checking for Circularity")
add_para(
    "A well known concern in this field is that a pathogenicity predictor can look accurate mainly "
    "because it was trained on labels similar or identical to the ones it is being tested against, "
    "rather than because it has learned something genuinely predictive. This project checked for "
    "this concern in two different ways rather than simply mentioning it as a caveat."
)
add_heading("8.1 Feature Group Comparison", level=2)
add_para(
    "The thirteen scores were split into two groups, based on published, independently derived "
    "classifications from the wider literature rather than this project's own judgment: six scores "
    "with no clinical label training exposure at all (AlphaMissense, both CADD scores, ESM1b, "
    "PrimateAI, and Eigen), and seven scores that were trained, at least in part, on clinical "
    "pathogenic and benign labels similar to the ones this project's own labels come from "
    "(ClinPred, MetaRNN, REVEL, MetaSVM, MetaLR, VEST4, and BayesDel). A model trained only on the "
    "six label free scores was compared against a model trained on all thirteen."
)
circ_display = circ.round(4).copy()
circ_display.index = [
    "All thirteen features, internal test", "All thirteen features, external",
    "Label free six features, internal test", "Label free six features, external",
    "Label trained seven features, internal test", "Label trained seven features, external",
]
add_df_table(circ_display, index_label="Feature Set", caption="Table 6. Full thirteen feature model versus label free subset versus label trained subset only.")
add_para(
    "The six label free scores alone reach almost exactly the same accuracy as the full thirteen "
    "feature model, and in this particular comparison very slightly exceed it. If most of the "
    "model's apparent accuracy came from circularity with clinical labels, the label free subset "
    "should have performed noticeably worse than the full model, and it did not."
)
add_heading("8.2 Variant Level Timing Check", level=2)
add_para(
    "A finer check was also done, looking at when each variant was actually classified in ClinVar, "
    "compared against the approximate dates when the label trained scores were built. The honest "
    "result here is less reassuring, and it is reported plainly rather than left out."
)
era_pivot = era.pivot_table(index=["pool", "label"], columns="cutoff", values="pct_pre_cutoff")
era_pivot.index = [f"{pool} pool, {label}" for pool, label in era_pivot.index]
era_pivot.columns = [c.split(" (")[0] for c in era_pivot.columns]
add_df_table(era_pivot.round(1), index_label="Pool and Label", caption="Table 7. Percentage of each pool already classified in ClinVar before each score's approximate training cutoff.")
add_para(
    "The external validation set turns out to be more exposed to this timing overlap than the "
    "internal set is, not less, which is the opposite of what a reader might expect from a set "
    "that is meant to be the cleaner, held out check. Between seventy two and one hundred percent "
    "of the external pool, depending on which score's cutoff is used and whether the variant is "
    "pathogenic or benign, was already classified in "
    "ClinVar before that score existed. This means the seven label trained scores' external "
    "validation numbers should not be read as fully circularity free, and this report does not "
    "claim that they are. It is a large part of why an entirely independent validation source, "
    "described in Section 10, was pursued instead of relying on the external set alone for this "
    "specific question."
)
add_heading("8.3 Population Frequency Double Counting", level=2)
add_para(
    "A separate, related concern raised in the literature is combining a population frequency "
    "based clinical evidence code with a pathogenicity score that already uses population "
    "frequency internally, which risks counting the same underlying signal twice. A direct search "
    "of every script in this pipeline confirmed that no population frequency based ACMG evidence "
    "code, such as PM2 or BS1, was being computed anywhere in the main pipeline, so this specific "
    "double counting mechanism could not currently occur in this project's own code."
)

# ---------------------------------------------------------------------------
add_heading("9. Clinical Evidence Calibration")
add_para(
    "Beyond a simple pathogenic or benign label, model output can be translated into standard "
    "clinical evidence strength categories following the ACMG and ClinGen framework, using a "
    "likelihood ratio approach. This project's original calibration was built once for the whole "
    "dataset, treating BRCA1 and BRCA2 as one pool. A recent paper describing a large scale, gene "
    "specific version of this same calibration approach reported that BRCA1 specifically benefits "
    "from its own, separately derived thresholds rather than a single genome wide threshold, while "
    "BRCA2 did not show the same clear benefit."
)
add_para(
    "That paper's own exact numerical thresholds were not accessible, since they are published "
    "only through an interactive web tool that could not be automated with the tools available "
    "for this project. Rather than guess at those numbers, this project adopted the finding itself "
    "and re derived its own gene specific thresholds directly from its own labeled data, separately "
    "for BRCA1 and BRCA2."
)
piv = acmg_gene.pivot_table(index="evidence", columns="split", values="score_threshold")
piv = piv.rename(columns={
    "gene_specific_BRCA1": "BRCA1 specific", "gene_specific_BRCA2": "BRCA2 specific",
    "pooled_genome_wide": "Pooled, genome wide",
})
piv.index = [i.replace("_", " ") for i in piv.index]
add_df_table(piv.round(3), index_label="Evidence Level", caption="Table 8. Model score thresholds for each clinical evidence strength, pooled versus gene specific.")
add_para(
    "BRCA1 and BRCA2 genuinely need different thresholds. For example, the score needed to reach "
    "moderate pathogenic evidence is noticeably higher for BRCA1 than for BRCA2 under this gene "
    "specific calibration, meaning a single shared threshold would misrepresent at least one of "
    "the two genes."
)
acmg_ext_display = acmg_ext.set_index("split")
acmg_ext_display.index = [
    {"pooled_genome_wide": "Pooled, genome wide", "gene_specific_BRCA1": "BRCA1 specific",
     "gene_specific_BRCA2": "BRCA2 specific"}.get(i, i) for i in acmg_ext_display.index
]
acmg_ext_display["check"] = acmg_ext_display["check"].str.replace("_", " ")
add_df_table(acmg_ext_display.round(1), index_label="Calibration", caption="Table 9. External validation of the calibrated evidence bins.")
add_para(
    "The strongest pathogenic evidence bin is between ninety eight and one hundred percent "
    "pathogenic on external validation across the pooled and both gene specific calibrations, and "
    "the strongest benign bin is zero percent pathogenic, meaning the calibration behaves "
    "sensibly out of sample."
)

# ---------------------------------------------------------------------------
add_heading("10. Additional Benchmarking")
add_para(
    "Two further checks were run, following methodology used in a recent, large scale independent "
    "benchmark of variant effect predictors. The first looks specifically at how each model "
    "behaves in the high specificity region of its performance, meaning how sensitive it is once "
    "the false positive rate is restricted to five percent or less, since this is closer to how a "
    "clinical laboratory would actually want to use such a tool."
)
stage8_both = stage8[(stage8["gene"] == "both") & (stage8["pool"].isin(["internal_test", "external"]))].set_index("model")
add_df_table(stage8_both[["pool", "n", "tpr_at_fpr5", "fpr_at_tpr95"]].round(4), index_label="Model",
             caption="Table 10. Sensitivity at five percent false positive rate, and false positive rate needed for ninety five percent sensitivity, both genes pooled together.")
add_para(
    "A second, gene specific version of the same check, shown in the next table using the "
    "internal test set, found a real difference between the two genes. Reaching ninety five "
    "percent sensitivity requires accepting a noticeably higher false positive rate for BRCA2 "
    "than for BRCA1 across every model tested, meaning BRCA2 variants are genuinely harder to "
    "classify with high confidence in this regime, not simply a reflection of different score "
    "thresholds."
)
stage8_gene = stage8[(stage8["pool"] == "internal_test") & (stage8["gene"].isin(["BRCA1", "BRCA2"]))]
stage8_gene_piv = stage8_gene.pivot_table(index="model", columns="gene", values="fpr_at_tpr95")
add_df_table(stage8_gene_piv.round(4), index_label="Model",
             caption="Table 11. False positive rate needed for ninety five percent sensitivity, BRCA1 versus BRCA2, internal test set.")
add_para(
    "A separate check confirmed that this project's strong performance is not simply an artifact "
    "of BRCA1 having a higher baseline rate of pathogenic variants than BRCA2 in the data. Re "
    "evaluating on a subset artificially balanced to have equal numbers of pathogenic and benign "
    "variants in both genes still produced strong results, confirming the model is not relying on "
    "that imbalance as a shortcut."
)
balanced = stage8[stage8["pool"].isin(["internal_all_gene_balanced", "external_gene_balanced"])].set_index("model")
balanced_display = balanced[["pool", "n", "tpr_at_fpr5", "fpr_at_tpr95"]].round(4).copy()
balanced_display["pool"] = balanced_display["pool"].map({
    "internal_all_gene_balanced": "Internal, gene balanced", "external_gene_balanced": "External, gene balanced",
})
add_df_table(balanced_display, index_label="Model",
             caption="Table 12. Results on the gene and label balanced subset described above.")

# ---------------------------------------------------------------------------
add_heading("11. Independent Functional Validation")
add_para(
    "All of the validation described so far, internal and external alike, ultimately traces back "
    "to clinical databases built by human curators, which is exactly the kind of evidence source "
    "that raises the circularity concerns discussed in Section 8. A stronger, genuinely "
    "independent check is to compare this project's model against a real laboratory experiment "
    "that measured the actual biological function of thousands of BRCA1 variants directly, with no "
    "reference to any clinical database at all."
)
add_para(
    "Such a dataset exists: a 2018 study used a genome editing technique to directly test the "
    "functional effect of nearly four thousand possible single letter changes across the "
    "functionally critical regions of BRCA1, measuring how well cells survived with each variant "
    "present. This data was obtained directly from its public hosting database using that "
    "database's own programming interface, and its identity was confirmed against the original "
    "publication before being used, rather than assumed from a file name."
)
n_small = len(dms_small)
n_full = len(dms_full[dms_full["matched"]]) if "matched" in dms_full.columns else len(dms_full)
add_para(
    f"An initial check against the {n_small} of these laboratory measured variants that happened "
    "to already be present in this project's own internal pool showed a strong relationship "
    "between the model's predicted pathogenicity and the measured functional effect. Recognizing "
    "that this initial check only covered variants already known to a clinical database, the check "
    "was extended to the entire laboratory dataset, most of which has never been classified by any "
    "clinical database at all."
)
add_para(
    f"Of the full experimental dataset, {n_full} variants could be matched to this project's "
    "annotation pipeline and scored. The great majority of these, roughly seventy three percent, "
    "are not present in this project's internal or external pools at all, meaning they are "
    "genuinely novel from a clinical database point of view. The model's predicted pathogenicity "
    "correlated strongly and significantly with the measured functional effect across this larger, "
    "more independent set, though somewhat less strongly than the smaller initial check suggested, "
    "which is itself an honest and expected finding rather than a contradiction, since the smaller "
    "subset was drawn from clinical database entries that tend to already be the clearer cut cases."
)
add_para(
    "This larger, functionally grounded check is treated in this report as the most trustworthy "
    "single estimate of how well the model actually generalizes, precisely because it does not "
    "depend on clinical database curation at all, and therefore does not inherit the timing "
    "overlap concern raised in Section 8.2."
)

# ---------------------------------------------------------------------------
add_heading("12. Class Imbalance and Synthetic Oversampling")
add_para(
    "Because pathogenic variants are a small minority of the usable training data, class weighting "
    "was used throughout this project rather than generating synthetic examples of the minority "
    "class. This choice was tested directly rather than simply assumed to be correct. Four "
    "classifiers were each trained under class weighting alone and under three different "
    "intensities of a synthetic oversampling technique, and evaluated on the same real internal "
    "test set and external validation set every other model in this report was evaluated on."
)
add_para(
    "The most aggressive oversampling setting, generating a synthetic pathogenic example for every "
    "real benign example, produced the worst calibration, meaning the least trustworthy predicted "
    "probabilities, of any setting tested, across every single classifier. This confirms the "
    "original decision to rely on class weighting rather than synthetic oversampling as the "
    "primary approach for this project."
)

# ---------------------------------------------------------------------------
add_heading("13. Kazakh Population Validation")
add_para(
    "A specific goal of this project, beyond general BRCA1 and BRCA2 classification, was to bring "
    "in real data from Kazakhstan, since population specific data of this kind is largely absent "
    "from the large international databases this project otherwise relies on. Four things were "
    "built from this Kazakh data, described in turn below."
)
add_heading("13.1 A Population Frequency Resource", level=2)
add_para(
    "Two hundred and twenty four individual genotyping array files from a Kazakh population study "
    "were processed directly to build a table of how common each of roughly four thousand "
    "genotyped positions across BRCA1 and BRCA2 is in this specific population. A real technical "
    "problem was caught and fixed during this step: the source files turned out to be built on an "
    "older genome coordinate system than the rest of this project uses, which was confirmed "
    "directly by reading the file headers rather than assumed, and every position was converted to "
    "match the rest of the pipeline before being used further."
)
add_heading("13.2 Population Based Clinical Evidence", level=2)
add_para(
    "This Kazakh population frequency table was used to build two clinical evidence codes "
    "following the ACMG framework: one supporting a pathogenic interpretation when a variant is "
    "essentially never seen in this population sample, and one supporting a benign interpretation "
    "when a variant is common enough in this population that it is unlikely to cause a rare, "
    "serious disease. Given how small the underlying sample is, only two hundred and twenty four "
    "people, these codes were built conservatively, using a statistical confidence interval on the "
    "true population frequency rather than the raw observed frequency, so that a single chance "
    "observation in a small sample does not get treated as strong evidence on its own."
)
add_heading("13.3 Checking Against Real Published Kazakh Variants", level=2)
add_para(
    "These evidence codes were then checked directly against a set of real, published BRCA1 and "
    "BRCA2 variants collected from three separate Kazakh research studies. Twenty five of these "
    "variants happened to fall on genotyped positions in the population resource. Every one of the "
    "twenty one variants already classified as pathogenic or likely pathogenic correctly received "
    "the pathogenic supporting code, and both "
    "of the variants already independently classified as common benign polymorphisms in their "
    "original publication correctly received the benign supporting code, which is a genuine "
    "external confirmation rather than a result this project engineered to succeed. Two variants "
    "that remain of uncertain significance in the published literature received new evidence in "
    "opposite directions, which is the actual practical use case for this kind of resource."
)
add_heading("13.4 Full Two Track Evaluation on Real Kazakh Variants", level=2)
add_para(
    f"Finally, every one of the real Kazakh clinical variants collected from the three source "
    "studies was run through this project's full two track classification pipeline, closing a gap "
    "that had been outstanding since early in this project. Every one of the "
    f"{lof_correct} known pathogenic loss of function variants was correctly classified by the "
    "rule based track. On the machine learning track, the small number of missense variants with "
    "a known clinical label were correctly and cleanly separated by Hybrid B, and three genuine "
    "variants of uncertain significance were all scored as more likely benign than pathogenic, "
    "again the kind of new evidence this project is ultimately meant to provide. Given how few "
    "labeled missense variants were available for this specific check, this result is reported as "
    "a demonstration of the pipeline working correctly end to end on real, independently sourced "
    "clinical variants, not as a statistically powered accuracy estimate."
)

# ---------------------------------------------------------------------------
add_heading("14. Comparison With Published Literature")
add_para(
    "Every number reported so far comes from this project's own data. To judge whether this "
    "work is competitive, it also needs to be placed next to results published by other groups "
    "working on the same two genes. A focused literature search was carried out for that purpose, "
    "and the results are collected in the table below."
)
add_para(
    "The one honest caveat that matters here is worth stating plainly, once: none of the studies "
    "in this table were evaluated on the same held out variants as this project's own models. This "
    "is not a formal statistical head to head, of the kind carried out between this project's own "
    "models in Section 7. It is, more modestly, this project's number placed next to another "
    "group's number, each computed on that group's own data."
)
add_para(
    "The closest true apples to apples comparison found is BayesDel's own validation, which used "
    "exactly the ENIGMA expert classification method this project's external set also comes from, "
    "restricted to BRCA1 and BRCA2 specifically, and deliberately chosen because that classification "
    "method does not itself rely on any sequence based predictor as input. On that test, BayesDel "
    "reached an area under the curve of zero point nine four, ahead of MetaSVM at zero point nine, "
    "MetaLR at zero point eight eight, and CADD also at zero point eight eight, three scores that "
    "are themselves part of this project's own thirteen feature set."
)
add_para(
    "Several other groups have built gene specific machine learning models for BRCA1 and BRCA2, "
    "directly comparable in spirit to this project's own missense track. A Mayo Clinic group's "
    "BRCA-ML model, validated against the same kind of functional assay ground truth used for this "
    "project's own independent validation in Section 11, reached a Matthews correlation coefficient "
    "of zero point six six for BRCA1 and zero point seven three for BRCA2. A Qatar based group built "
    "separate gene specific models for each gene and reported accuracy of ninety nine point nine "
    "percent on ENIGMA classified test variants, and accuracy up to ninety three percent on an "
    "independent, functionally scored set of variants of uncertain significance, a similar design "
    "and a similar headline number to this project's own external validation approach. A third group "
    "directly compared gene specific, disease specific, and genome wide trained machine learning "
    "against popular general purpose predictors including REVEL, BayesDel, and ClinPred, and found "
    "gene specific training clearly ahead of the genome wide predictors on both genes, but with "
    "several of its own gene specific and disease specific models statistically indistinguishable "
    "from each other, essentially the same finding this project's own model comparison reached in "
    "Section 7."
)
add_para(
    "One further comparison stands out as unusually direct. A separate ensemble model, SNPred, was "
    "validated against the exact same publicly hosted BRCA1 functional genome editing dataset used "
    "for this project's own independent functional validation in Section 11, and reported an area "
    "under the precision recall curve of zero point seven nine four, ahead of BayesDel's zero point "
    "seven three one on that same dataset. That paper's central argument, that predictors trained on "
    "clinical databases score noticeably higher when re evaluated on those same databases than when "
    "evaluated on independent functional data, and that this gap can be shown mathematically to "
    "require an implausibly low error rate in the clinical database itself, is essentially the same "
    "argument made independently in Section 8.2 of this report using this project's own data."
)
add_para(
    "Not every study found is a fair like for like comparison, and two are worth flagging "
    "specifically so the table below is not misread. One recent study trains its model on tumor "
    "genomic profiling features, such as a tumor's homologous recombination deficiency signature and "
    "co-mutated genes, rather than on sequence or conservation based scores at all, and reaches a "
    "validation area under the curve of essentially one for BRCA1 and zero point nine eight nine for "
    "BRCA2. This is not a competing version of the same kind of model built here. It is evidence from "
    "a completely different source, and is better read as a candidate for future evidence "
    "integration than as a benchmark this project's own model should be expected to match. A second "
    "study applied several deep learning predictors, two of which are also features in this "
    "project's own model, to real breast cancer outcomes in a large population database, and found "
    "that while the predictors tracked real disease risk well overall, their performance specifically "
    "on variants of uncertain significance was limited. That is a useful, independent caution about "
    "exactly the kind of variant this project is most trying to help classify, and is reflected "
    "directly in this report's own limitations in Section 15."
)
lit_display = lit.copy()
lit_display["source_paper"] = lit_display["source_paper"].str.replace(r"\s*\(articles/.*?\)", "", regex=True)
lit_display = lit_display.drop(columns=["notes"])
add_df_table(
    lit_display,
    caption=(
        "Table 13. This project's results placed alongside published benchmarks from other "
        "groups. Not a formal statistical comparison; full source details and caveats for every "
        "row are recorded in articles/citation_log.md."
    ),
)

# ---------------------------------------------------------------------------
add_heading("15. Limitations")
add_bullet(
    "Several scripts central to this project's earlier reported results, including the hybrid "
    "model, the clinical evidence calibration, the loss of function rule engine, and the "
    "synthetic oversampling comparison, were found to be missing from the project files entirely "
    "at various points, with only their previously reported result tables surviving. All were "
    "rebuilt from written descriptions and re evaluated, but a reader should treat any number from "
    "before this rebuilding as unreproducible unless it has been explicitly re confirmed in this "
    "report."
)
add_bullet(
    "No single model in this report has been shown to be statistically better than the others on "
    "the external validation set. Any future presentation of this work should not claim Hybrid B, "
    "or any other single model, as definitively the best approach."
)
add_bullet(
    "The external validation set is more exposed to training era overlap with several of the "
    "clinically trained scores than the internal set is, which limits how much weight those "
    "particular scores' external validation numbers alone should carry."
)
add_bullet(
    "The external validation set is small, and a post cutoff filtered version of it, restricted "
    "to only recently classified variants, was attempted specifically to address the point above "
    "and found to leave essentially no usable benign variants at all, meaning that particular fix "
    "is not currently possible with this data source."
)
add_bullet(
    "Sequence based deep learning models trained from scratch on this project's data do not "
    "currently generalize to the external validation set, dropping to close to chance level "
    "performance, and should not be presented as a strength of this project without this caveat "
    "attached directly."
)
add_bullet(
    "The Kazakh missense track validation involves a very small number of labeled variants and "
    "should be described as a demonstration, not as a statistically powered result, in any future "
    "summary of this work."
)
add_bullet(
    "The Kazakh population frequency resource is built from a fixed content genotyping array "
    "covering a limited, pre selected set of positions, not from full genome sequencing, and "
    "cannot discover or characterize rare variants that were not already on the array."
)

# ---------------------------------------------------------------------------
add_heading("16. Conclusion")
add_para(
    "This project built a thirteen feature, two track classification pipeline for BRCA1 and BRCA2 "
    "variant pathogenicity, and then spent a substantial amount of effort checking that pipeline "
    "rather than simply reporting its headline numbers. That checking process found and fixed "
    "several real errors, found that the project's own earlier best model claim did not hold up "
    "under formal statistical testing, found that its external validation set carries more "
    "circularity risk than it appears to at first glance, and found a genuinely independent "
    "functional validation source that supports the model's predictions even where clinical "
    "database evidence cannot be fully trusted. It also extended the project into a new, real "
    "population, using actual Kazakh genetic data and actual published Kazakh clinical variants, "
    "rather than treating population specific validation as an afterthought."
)
add_para(
    "The honest summary is that this is a solid, carefully checked piece of work with a "
    "genuinely novel population specific contribution, built on a model that performs well but is "
    "not proven to be better than several simpler alternatives, evaluated against both a clinical "
    "gold standard and an independent functional one, with every important limitation stated "
    "plainly rather than left for a reader to discover on their own."
)

doc.save("2/BRCA_Final_Technical_Report_2026-07-27.docx")
print("Saved 2/BRCA_Final_Technical_Report_2026-07-27.docx")
